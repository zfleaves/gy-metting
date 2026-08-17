"""
文档解析器 (DESIGN.md §3.2)

统一接口 parse(file_path) → str，支持 docx / pdf / txt / md。
"""

from pathlib import Path
from typing import Optional


def parse(file_path: str) -> str:
    """
    解析文档文件，返回纯文本内容。

    Args:
        file_path: 文档文件路径

    Returns:
        解析后的纯文本

    Raises:
        ValueError: 不支持的格式
        FileNotFoundError: 文件不存在
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = path.suffix.lower()
    if ext == ".docx":
        return _parse_docx(path)
    elif ext == ".pdf":
        return _parse_pdf(path)
    elif ext in (".txt", ".md", ".markdown"):
        return _parse_text(path)
    else:
        raise ValueError(f"不支持的文档格式: {ext}，支持 docx/pdf/txt/md")


def detect_format(file_path: str) -> str:
    """检测文档格式，返回小写扩展名"""
    return Path(file_path).suffix.lower().lstrip(".")


def _parse_docx(path: Path) -> str:
    """解析 docx 文件"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    if not paragraphs:
        # 尝试从表格中提取
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def _parse_pdf(path: Path) -> str:
    """解析 pdf 文件"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF 未安装，请运行: pip install PyMuPDF")

    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


def _parse_text(path: Path) -> str:
    """解析纯文本/markdown 文件"""
    return path.read_text(encoding="utf-8")